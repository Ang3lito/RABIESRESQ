"""One-off generator for RabiesResQ Installation and User Guide (.docx)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_h(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_p(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_num_steps(doc: Document, items: list[str]):
    for t in items:
        p = doc.add_paragraph(t, style="List Number")
        p.paragraph_format.space_after = Pt(3)


def add_bullets_black(doc: Document, items: list[str]):
    """Bullet list using ⚫ to match organization style guides."""
    for t in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.add_run("⚫ ")
        p.add_run(t)


def main():
    out = Path(__file__).resolve().parent.parent / "RabiesResQ_Installation_and_User_Guide.docx"
    doc = Document()

    t = doc.add_heading("RabiesResQ", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = doc.add_paragraph("Installation Guide and User Guide")
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.runs[0].bold = True
    doc.add_paragraph()

    add_p(
        doc,
        "This manual describes how to access and use RabiesResQ, and includes technical "
        "notes for system administrators who host the application.",
    )

    # --- Part 1 ---
    add_h(doc, "Part 1 — Installation Guide", 1)

    add_h(doc, "What You Need", 2)
    add_bullets_black(
        doc,
        [
            "A device (computer, laptop, tablet, or smartphone)",
            "Stable internet connection",
            "A modern web browser (Chrome, Firefox, Edge, or Safari)",
        ],
    )

    add_h(doc, "How to Access", 2)
    add_num_steps(
        doc,
        [
            "Open your web browser.",
            "Go to: https://yourusername.pythonanywhere.com/ (replace yourusername with your PythonAnywhere username, or use the exact URL your organization provides if you have a custom domain).",
            "The system will automatically load and adjust to your screen size.",
        ],
    )
    add_p(
        doc,
        "If you do not know the correct address, ask your clinic or system administrator.",
    )

    add_h(doc, "For system administrators: hosting and deployment", 2)
    add_p(
        doc,
        "RabiesResQ is a browser-based application for rabies post-exposure workflows: "
        "patients, clinic personnel, and system administrators use roles to access cases, "
        "appointments, vaccination records, and reporting. The application is built with "
        "Python Flask, served via WSGI, and uses SQLite. The following steps describe "
        "deployment on PythonAnywhere.",
    )

    add_h(doc, "1. Prepare the environment", 3)
    add_num_steps(
        doc,
        [
            "On PythonAnywhere, open a Bash console. On a local machine, use a terminal with Python 3 and pip available.",
            "Clone or upload the RabiesResQ project into a directory under your home folder (for example /home/yourusername/RABIESRESQ).",
            "Create a virtual environment and activate it. On PythonAnywhere, use the same Python version you will select for the web app.",
            "Install dependencies: pip install -r requirements.txt",
        ],
    )

    add_h(doc, "2. Configure environment variables", 3)
    add_num_steps(
        doc,
        [
            "Create a .env file in the project root (the folder that contains app.py). The application loads this file automatically.",
            "Set SECRET_KEY to a strong random value (required for sessions and security tokens).",
            "Optionally set DATABASE to the full absolute path of your SQLite file; otherwise the app uses the default under the Flask instance folder (ensure this path is writable on PythonAnywhere).",
            "For outbound email, set MAIL_USERNAME and MAIL_PASSWORD when using Gmail SMTP as implemented in the application.",
            "Alternatively, you may define the same variables in the PythonAnywhere Web tab under Environment variables, if you prefer not to rely on a .env file on the server.",
        ],
    )

    add_h(doc, "3. Initialize the database", 3)
    add_num_steps(
        doc,
        [
            "Apply schema.sql to your SQLite database using the sqlite3 command line or another SQLite tool, targeting the same path the application will use (including under instance/ if you use the default).",
            "Ensure at least one clinic exists and provision administrative accounts before go-live.",
        ],
    )

    add_h(doc, "4. Create initial clinic and privileged accounts", 3)
    add_p(
        doc,
        "With the virtual environment activated, change to the project directory and run Flask CLI commands "
        "(flask --app app:create_app …). On PythonAnywhere, run these from a Bash console, not from the PythonAnywhere Python REPL.",
    )
    add_num_steps(
        doc,
        [
            "create-clinic — name and optional address.",
            "create-admin — username, email, password, employee ID, optional names.",
            "create-staff — username, email, password, clinic ID, employee ID, title (Doctor or Nurse), optional license and names.",
        ],
    )
    add_p(doc, "Additional staff may be created from the administrator user interface after deployment.")

    add_h(doc, "5. Production deployment on PythonAnywhere", 3)
    add_p(
        doc,
        "These steps follow the usual PythonAnywhere workflow for a Flask WSGI application. "
        "Exact labels may vary slightly as the PythonAnywhere dashboard is updated; use their documentation for the current Web configuration screens.",
    )
    add_num_steps(
        doc,
        [
            "In the PythonAnywhere Dashboard, open the Web tab and create a web app if you have not already. Choose Manual configuration and select the Python version that matches your virtual environment.",
            "Set the virtual environment path on the Web tab to the virtualenv you created for RabiesResQ.",
            "Set the working directory (if the Web UI offers it) to your project folder, or rely on sys.path in the WSGI file as below.",
            "Edit the WSGI configuration file that PythonAnywhere provides for your site. Add your project directory to sys.path at the top of the file, then import the WSGI application object. For example, if the project lives at /home/yourusername/RABIESRESQ, add that path to sys.path, then use: from wsgi import application (the repository includes wsgi.py exposing application = create_app()).",
            "Ensure the WSGI file does not call app.run(); the web server invokes application directly.",
            "Optional: map the URL path /static/ to your project’s static/ folder in the Static files section of the Web tab so the web server can serve static assets efficiently.",
            "Click Reload to load configuration changes. Visit your site URL and confirm the login page loads.",
            "After each code update (for example git pull), reload the web app again so changes take effect.",
        ],
    )

    add_h(doc, "6. Local development (optional)", 3)
    add_num_steps(
        doc,
        [
            "For quick testing on your own computer, with .env and the database in place: flask --app app:create_app run",
            "Optionally bind to all interfaces: --host 0.0.0.0 --port 5000",
            "Do not use the Flask development server for production; use the PythonAnywhere WSGI deployment above for live hosting.",
        ],
    )
    add_p(
        doc,
        "Email for password reset requires SMTP configuration (MAIL_USERNAME and MAIL_PASSWORD). "
        "On PythonAnywhere, confirm your account tier allows outbound SMTP; if email is not configured, "
        "reset codes will not be delivered to users.",
    )

    # --- Part 2 ---
    add_h(doc, "Part 2 — User Guide", 1)

    add_h(doc, "1. General account access", 2)
    add_num_steps(
        doc,
        [
            "Each person uses one account identified by email and password.",
            "Accounts have a role that controls available screens and actions.",
            "Inactive accounts cannot sign in; contact an administrator.",
            "If a password change is required, patients and clinic personnel must set a new password before using the rest of the application.",
        ],
    )

    add_h(doc, "2. Login", 2)
    add_num_steps(
        doc,
        [
            "Open the application URL.",
            "Enter email and password and submit.",
            "You are taken to the home area for your role; first-time patients may complete onboarding first.",
        ],
    )

    add_h(doc, "3. Registration and account creation", 2)
    add_h(doc, "3.1 Patient self-registration", 3)
    add_num_steps(
        doc,
        [
            "From the login page, open registration.",
            "Complete username, email, password, confirm password, and optional profile fields.",
            "Submit; the system creates a patient account and signs you in.",
        ],
    )
    add_h(doc, "3.2 Staff and administrator accounts", 3)
    add_p(
        doc,
        "Clinic personnel and system administrators are not created via public self-registration. "
        "They are created by a system administrator (UI and/or CLI) using your secure process.",
    )

    add_h(doc, "4. Password recovery", 2)
    add_num_steps(
        doc,
        [
            "Open forgot password from the login page.",
            "Enter your email and request a code.",
            "If an account exists, a six-digit verification code is sent when email is configured.",
            "Enter the code; codes expire after a short period.",
            "Complete password reset with a new password and confirmation (minimum eight characters).",
            "Sign in with the new password.",
        ],
    )

    add_h(doc, "5. Dashboard overview (by role)", 2)
    add_num_steps(
        doc,
        [
            "Patients: dashboard summarizes cases and appointments (including dependents where applicable).",
            "Clinic personnel: operational snapshot for your clinic with indicators where shown.",
            "System administrators: cases, appointments, reporting, users, Session Logs, and settings; login may land on reporting/analytics.",
        ],
    )

    add_h(doc, "6. Navigation overview", 2)
    add_num_steps(
        doc,
        [
            "Use sidebar navigation; on small screens open the menu first.",
            "Profile: use the sidebar profile area (patient/staff) or settings entry (admin).",
            "Logout is available from the sidebar on authenticated pages.",
        ],
    )

    add_h(doc, "7. Managing records and data", 2)
    add_h(doc, "7.1 Patients", 3)
    add_num_steps(
        doc,
        [
            "Complete first-time onboarding when prompted.",
            "Start a new appointment using the pre-screening flow when available.",
            "Open appointments to view details, request changes, cancel, or hide items as offered.",
            "Review vaccinations and vaccination card views; download PDF where offered.",
            "Use availability features per your deployment.",
            "Update your profile.",
        ],
    )
    add_h(doc, "7.2 Clinic personnel", 3)
    add_num_steps(
        doc,
        [
            "Cases: list, open, create, edit; notes, WHO category handling, complete or remove cases per policy.",
            "Patients: register new patient accounts; link workflows to the correct person.",
            "Appointments: approve, edit, or remove as policy allows.",
            "Availability: maintain clinic slots.",
            "Vaccinations: work from the vaccinations area; export CSV or PDF when needed.",
            "Operations: date-bounded summaries for your clinic.",
            "Records: case record PDFs and case/vaccination exports where available.",
            "Profile: maintain staff profile.",
        ],
    )
    add_h(doc, "7.3 System administrators", 3)
    add_num_steps(
        doc,
        [
            "Cases and appointments: browse and monitor with filters as provided.",
            "Reporting: overview, clinic, and insights tabs; CSV/PDF exports where exposed.",
            "Users: activate or deactivate accounts (with safeguards); create new staff.",
            "Session Logs: review sign-in and sign-out history.",
            "Settings: clinic and system options on the settings screen.",
            "Clinic exports: CSV or PDF when offered.",
        ],
    )

    add_h(doc, "8. Viewing reports or summaries", 2)
    add_num_steps(
        doc,
        [
            "Clinic personnel: Operations page for period summaries.",
            "System administrators: Reporting for analytics and optional forensic or insights exports.",
            "Exports are started from the relevant screens.",
        ],
    )

    add_h(doc, "9. Notifications", 2)
    add_num_steps(
        doc,
        [
            "In-app notifications may appear on dashboards.",
            "Optional notification sounds may play when enabled and permitted by the browser.",
            "Email is used for password reset; not all in-app notices are emailed.",
        ],
    )

    add_h(doc, "10. Profile management", 2)
    add_num_steps(
        doc,
        [
            "Patients and clinic personnel: profile pages from the sidebar.",
            "Administrators: settings and related screens for organization configuration.",
        ],
    )

    add_h(doc, "11. Other common actions", 2)
    add_num_steps(
        doc,
        [
            "Sign out on shared computers.",
            "Use Help on the patient interface where available.",
            "Follow minimum password length and organizational data-handling rules.",
        ],
    )

    # Role sections
    add_h(doc, "Role-based user guide sections", 1)

    add_h(doc, "A. Patient", 2)
    add_p(doc, "Typical tasks:")
    add_num_steps(
        doc,
        [
            "Register and complete onboarding.",
            "View the dashboard for appointments and case-related information.",
            "Start a new appointment using pre-screening when that is how the clinic accepts requests.",
            "Open appointments; edit, cancel, or hide as allowed.",
            "Browse vaccinations and vaccination cards; download PDFs when available.",
            "Check availability to align with clinic slots.",
            "Update profile and review Help.",
        ],
    )
    add_p(doc, "Note: Self-registration creates a patient account only.")

    add_h(doc, "B. Clinic personnel (doctors and nurses)", 2)
    add_p(doc, "Typical tasks:")
    add_num_steps(
        doc,
        [
            "Review the dashboard for work queues.",
            "Manage cases: new patient flows, add existing record, notes, WHO category, complete or delete cases.",
            "Handle appointments: approve, edit, or remove.",
            "Maintain availability.",
            "Work in vaccinations and use exports.",
            "Use Operations for period summaries.",
            "Use search or lists to find patients and cases.",
            "Update staff profile.",
        ],
    )
    add_p(doc, "Staff accounts are tied to a clinic, title (Doctor or Nurse), and unique identifiers.")

    add_h(doc, "C. System administrator", 2)
    add_p(doc, "Typical tasks:")
    add_num_steps(
        doc,
        [
            "Oversee cases and appointments.",
            "Use reporting and downloads where shown.",
            "Manage users and onboard new clinic staff.",
            "Review Session Logs.",
            "Configure settings.",
            "Use CLI commands for initial clinic creation when IT procedures allow.",
        ],
    )
    add_p(doc, "Only designated personnel should receive this role.")

    doc.add_paragraph()
    end = doc.add_paragraph("End of manual")
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end.runs[0].italic = True

    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
